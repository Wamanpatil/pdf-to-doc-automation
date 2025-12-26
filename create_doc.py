from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def write_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text if text else "")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.bold = bold

# ======================
# TITLE (NO EXTRA SPACE)
# ======================
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FORM ‘A’\nMEDIATION APPLICATION FORM\n[REFER RULE 3(1)]")
r.font.name = "Times New Roman"
r.font.size = Pt(14)
r.font.bold = True

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Mumbai District Legal Services Authority\nCity Civil Court, Mumbai")
r.font.name = "Times New Roman"
r.font.size = Pt(12)
r.font.bold = True

# ======================
# DETAILS OF PARTIES TABLE
# ======================
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"

write_cell(table.rows[0].cells[0], "DETAILS OF PARTIES", True)
write_cell(table.rows[0].cells[1], "")

rows = [
    ("1. Name of Applicant", "{{client_name}}"),
    ("Registered Address", "{{branch_address}}"),
    ("Correspondence Branch Address", "{{branch_address}}"),
    ("Telephone No.", "{{mobile}}"),
    ("Mobile No.", ""),
    ("Email ID", "info@kslegal.co.in"),
    ("2. Name of Opposite Party", "{{customer_name}}"),
    ("Registered Address", "__________"),
    ("Correspondence Address", "__________"),
    ("Telephone No.", ""),
    ("Mobile No.", ""),
    ("Email ID", "")
]

for left, right in rows:
    r = table.add_row().cells
    write_cell(r[0], left, left.startswith(("1.", "2.")))
    write_cell(r[1], right)

# ======================
# DISPUTE TABLE
# ======================
dispute = doc.add_table(rows=3, cols=1)
dispute.style = "Table Grid"

write_cell(dispute.rows[0].cells[0], "DETAILS OF DISPUTE", True)
write_cell(
    dispute.rows[1].cells[0],
    "THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018",
    True
)
write_cell(
    dispute.rows[2].cells[0],
    "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
)

# ======================
# SAVE (CHANGE NAME TO AVOID LOCK)
# ======================
doc.save("output/mediation_application_form_final.docx")
print("MS Word document created successfully.")
