from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

app = Flask(__name__)
OUTPUT_FILE = "output/mediation_application_form.docx"

def write_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text if text else "")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.bold = bold

def generate_doc():
    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FORM ‘A’\nMEDIATION APPLICATION FORM\n[REFER RULE 3(1)]\n")
    r.font.size = Pt(14)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mumbai District Legal Services Authority\nCity Civil Court, Mumbai")
    r.font.size = Pt(12)
    r.font.bold = True

    # Table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    write_cell(table.rows[0].cells[0], "DETAILS OF PARTIES", True)
    write_cell(table.rows[0].cells[1], "")

    rows = [
        ("1. Name of Applicant", "{{client_name}}"),
        ("Registered Address", "{{branch_address}}"),
        ("Correspondence Branch Address", "{{branch_address}}"),
        ("Telephone No.", "{{mobile}}"),
        ("Mobile No.", ""),
        ("Email ID", "info@kslegal.co.in"),
        ("2. Name of Opposite Party", "{{customer_name}}"),
        ("Registered Address", "__________"),
        ("Correspondence Address", "__________"),
        ("Telephone No.", ""),
        ("Mobile No.", ""),
        ("Email ID", "")
    ]

    for left, right in rows:
        r = table.add_row().cells
        write_cell(r[0], left, left.startswith(("1.", "2.")))
        write_cell(r[1], right)

    os.makedirs("output", exist_ok=True)
    doc.save(OUTPUT_FILE)

@app.route("/", methods=["GET"])
def index():
    return """
    <h2>PDF to MS Word Automation</h2>
    <form action="/generate" method="post">
        <button type="submit">Generate Word Document</button>
    </form>
    """

@app.route("/generate", methods=["POST"])
def generate():
    generate_doc()
    return send_file(OUTPUT_FILE, as_attachment=True)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)